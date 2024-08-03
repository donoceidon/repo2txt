#!/usr/bin/env python

import os
import argparse
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def load_config(file_path):
    """
    Load configuration from a JSON file.

    Args:
        file_path (str): The path to the JSON configuration file.

    Returns:
        dict: A dictionary containing the configuration settings.
    """
    with open(file_path, 'r') as file:
        return json.load(file)

# Usage
config = load_config('config.json')
IMAGE_EXTENSIONS = config["image_extensions"]
VIDEO_EXTENSIONS = config["video_extensions"]
AUDIO_EXTENSIONS = config["audio_extensions"]
DOCUMENT_EXTENSIONS = config["document_extensions"]
EXECUTABLE_EXTENSIONS = config["executable_extensions"]
SETTINGS_EXTENSIONS = config["settings_extensions"]
ADDITIONAL_IGNORE_TYPES = config["additional_ignore_types"]
DEFAULT_OUTPUT_FILE = config["default_output_file"]

def parse_args():
    """
    Parse command-line arguments for the script.

    Returns:
        argparse.Namespace: An object containing the parsed command-line arguments.
    """

    parser = argparse.ArgumentParser(
        description='Document the structure of a GitHub repository.',
        epilog=('Example usage:\n'
                'Direct script invocation:\n'
                '  python repo2txt.py -r /path/to/repo -o output.txt  # Save as text\n'
                '  python repo2txt.py -r /path/to/repo -o report.docx  # Save as DOCX\n'
                'When installed with pip as a command-line tool:\n'
                '  repo2txt -r /path/to/repo -o output.txt  # Save as text\n'
                '  repo2txt -r /path/to/repo -o report.docx  # Save as DOCX\n\n'
                'Note: Specify the output file format by choosing the appropriate file extension (.txt or .docx).'),
        formatter_class=argparse.RawDescriptionHelpFormatter
    )

    parser.add_argument('-r', '--repo_path', default=os.getcwd(),
                        help='Path to the directory to process (ie., cloned repo). if no path is specified defaults to the current directory.')
    parser.add_argument('-o', '--output_file', default=DEFAULT_OUTPUT_FILE,
                        help='Name for the output text file. Defaults to "output.txt".')
    parser.add_argument('--ignore-files', nargs='*', default=[],
                        help='List of file names to ignore. Omit this argument to ignore no file names.')
    parser.add_argument('--ignore-types', nargs='*', default=IMAGE_EXTENSIONS + VIDEO_EXTENSIONS + AUDIO_EXTENSIONS + DOCUMENT_EXTENSIONS + EXECUTABLE_EXTENSIONS,
                        help='List of file extensions to ignore. Defaults to list in config.json. Omit this argument to ignore no types.')
    parser.add_argument('--exclude-dir', nargs='*', default=[],
                        help='List of directory names to exclude or "none" for no directories.')
    parser.add_argument('--ignore-settings', action='store_true',
                        help='Flag to ignore common settings files.')
    parser.add_argument('--include-dir', nargs='?', default=None,
                        help='Specific directory to include. Only contents of this directory will be documented.')

    return parser.parse_args()


def should_ignore(item, args, output_file_path):

    """
    Determine if a given item should be ignored based on the script's arguments.

    Args:
        item (str): The path of the item (file or directory) to check.
        args (argparse.Namespace): Parsed command-line arguments.
        output_file_path (str): The path of the output file being written to.

    Returns:
        bool: True if the item should be ignored, False otherwise.
    """

    item_name = os.path.basename(item)
    file_ext = os.path.splitext(item_name)[1].lower()

    # Ensure the comparison is between path strings
    if os.path.abspath(item) == os.path.abspath(output_file_path):
        return True

    # Adjust logic to handle hidden files and directories correctly
    if item_name.startswith('.'):
        return True  # Ignore all hidden files and directories

    if os.path.isdir(item) and args.exclude_dir and item_name in args.exclude_dir:
        return True

    if args.include_dir and not os.path.abspath(item).startswith(os.path.abspath(args.include_dir)):
        return True

    if os.path.isfile(item) and (item_name in args.ignore_files or file_ext in args.ignore_types):
        return True

    if args.ignore_settings and file_ext in SETTINGS_EXTENSIONS:
        return True

    return False

def write_tree(dir_path, output_file, args, prefix="", is_last=True, is_root=True):
    """
    Recursively write the directory tree to the output file, including the root directory name.

    Args:
        dir_path (str): The path of the directory to document.
        output_file (file object): The file object to write to.
        args (argparse.Namespace): Parsed command-line arguments.
        prefix (str): Prefix string for line indentation and structure. Defaults to "".
        is_last (bool): Flag to indicate if the item is the last in its level. Defaults to True.
        is_root (bool): Flag to indicate if the current directory is the root. Defaults to True.
    """

    if is_root:
        output_file.write(f"{os.path.basename(dir_path)}/\n")
        is_root = False

    items = os.listdir(dir_path)
    items.sort()  # Optional: Sort the items for consistent order
    num_items = len(items)

    for index, item in enumerate(items):
        item_path = os.path.join(dir_path, item)

        if should_ignore(item_path, args, args.output_file):
            continue

        is_last_item = (index == num_items - 1)
        new_prefix = "└── " if is_last_item else "├── "
        child_prefix = "    " if is_last_item else "│   "

        output_file.write(f"{prefix}{new_prefix}{os.path.basename(item)}\n")

        if os.path.isdir(item_path):
            next_prefix = prefix + child_prefix
            write_tree(item_path, output_file, args, next_prefix, is_last_item, is_root=False)


def write_file_content(file_path, output_file, depth):
    """
    Write the contents of a given file to the output file with proper indentation.

    Args:
        file_path (str): Path of the file to read.
        output_file (file object): The file object to write the contents to.
        depth (int): Current depth in the directory tree for indentation.
    """
    indentation = '  ' * depth
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            for line in file:
                # Indent each line of the file content
                output_file.write(f"{indentation}{line}")
    except Exception as e:
        output_file.write(f"{indentation}Error reading file: {e}\n")


def write_tree_docx(dir_path, doc, args, output_file_path, prefix="", is_last=True, is_root=True):
    """
    Recursively create a document structure of the directory tree in a DOCX file, including the root directory name.

    Args:
        dir_path (str): The path of the directory to document.
        doc (Document): The DOCX document object to write to.
        args (argparse.Namespace): Parsed command-line arguments.
        output_file_path (str): The path of the output DOCX file being written to.
        prefix (str): Prefix string for line indentation and structure. Defaults to "".
        is_last (bool): Flag to indicate if the item is the last in its level. Defaults to True.
        is_root (bool): Flag to indicate if the current directory is the root. Defaults to True.
    """

    if is_root:
        root_paragraph = doc.add_paragraph()
        root_paragraph.add_run(f"{os.path.basename(dir_path)}/")
        is_root = False

    items = os.listdir(dir_path)
    items.sort()  # Optional: Sort the items for consistent order
    num_items = len(items)

    for index, item in enumerate(items):
        item_path = os.path.join(dir_path, item)

        if should_ignore(item_path, args, output_file_path):
            continue

        is_last_item = (index == num_items - 1)
        new_prefix = "└── " if is_last_item else "├── "
        child_prefix = "    " if is_last_item else "│   "

        # Add the directory or file entry
        tree_paragraph = doc.add_paragraph()
        tree_paragraph.add_run(f"{prefix}{new_prefix}{os.path.basename(item)}")

        if os.path.isdir(item_path):
            next_prefix = prefix + child_prefix
            write_tree_docx(item_path, doc, args, output_file_path, next_prefix, is_last_item, is_root=False)


def write_file_content_docx(file_path, doc):

    """
    Write the contents of a given file to a DOCX document.

    Args:
        file_path (str): Path of the file to read.
        doc (Document): The DOCX document object to write the contents to.

    This function reads the contents of 'file_path' and writes them to 'doc'.
    If an error occurs during reading, it adds an error message to 'doc'.
    """

    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            contents = file.read()
            paragraph = doc.add_paragraph(contents)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    except Exception as e:
        error_paragraph = doc.add_paragraph(f"Error reading file: {e}")
        error_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT



def write_file_contents_in_order(dir_path, output_file, args, depth=0):
    """
    Recursively document the contents of files in the order they appear in the directory tree.

    Args:
        dir_path (str): The path of the directory to start documenting from.
        output_file (file object): The file object to write the contents to.
        args (argparse.Namespace): Parsed command-line arguments.
        depth (int): Current depth in the directory tree. Defaults to 0.
    """
    items = sorted(item for item in os.listdir(dir_path) if not should_ignore(os.path.join(dir_path, item), args, args.output_file))

    for item in items:
        item_path = os.path.join(dir_path, item)
        relative_path = os.path.relpath(item_path, start=args.repo_path)

        if os.path.isdir(item_path):
            write_file_contents_in_order(item_path, output_file, args, depth + 1)
        elif os.path.isfile(item_path):
            output_file.write('  ' * depth + f"[File Begins] {relative_path}\n")
            write_file_content(item_path, output_file, depth)
            output_file.write('\n' + '  ' * depth + f"[File Ends] {relative_path}\n\n")


def write_file_contents_in_order_docx(dir_path, doc, args, depth=0):
    """
    Recursively document the contents of files in a DOCX document in the order they appear in the directory tree.

    Args:
        dir_path (str): The path of the directory to start documenting from.
        doc (Document): The DOCX document object to write the contents to.
        args (argparse.Namespace): Parsed command-line arguments.
        depth (int): Current depth in the directory tree. Defaults to 0.
    """
    items = sorted(item for item in os.listdir(dir_path) if not should_ignore(os.path.join(dir_path, item), args, args.output_file))

    for item in items:
        item_path = os.path.join(dir_path, item)
        relative_path = os.path.relpath(item_path, start=args.repo_path)

        if os.path.isdir(item_path):
            write_file_contents_in_order_docx(item_path, doc, args, depth + 1)
        elif os.path.isfile(item_path):
            doc.add_heading(f"[File Begins] {relative_path}", level=3)
            write_file_content_docx(item_path, doc)
            doc.add_heading(f"[File Ends] {relative_path}", level=3)

def main():

    """
    Main function to execute the script logic.
    """
    args = parse_args()

    # Convert 'none' keyword to empty list
    args.ignore_files = [] if args.ignore_files == ['none'] else args.ignore_files
    args.ignore_types = [] if args.ignore_types == ['none'] else IMAGE_EXTENSIONS + VIDEO_EXTENSIONS + AUDIO_EXTENSIONS + DOCUMENT_EXTENSIONS + EXECUTABLE_EXTENSIONS + ADDITIONAL_IGNORE_TYPES
    args.exclude_dir = [] if args.exclude_dir == ['none'] else args.exclude_dir

    # Check if the provided directory path is valid
    if not os.path.isdir(args.repo_path):
        print(f"Error: The specified directory does not exist, path is wrong or is not a directory: {args.repo_path}")
        return  # Exit the script

    if args.output_file.endswith('.docx'):
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(11)
        
        doc.add_heading("Repository Documentation", level=1)
        doc.add_paragraph(
        "This document provides a comprehensive overview of the repository's structure and contents."
        "The first section, titled 'Directory/File Tree', displays the repository's hierarchy in a tree format."
        "In this section, directories and files are listed using tree branches to indicate their structure and relationships."
        "Following the tree representation, the 'File Content' section details the contents of each file in the repository."
        "Each file's content is introduced with a '[File Begins]' marker followed by the file's relative path,"
        "and the content is displayed verbatim. The end of each file's content is marked with a '[File Ends]' marker."
        "This format ensures a clear and orderly presentation of both the structure and the detailed contents of the repository.\n\n"
        )
        doc.add_heading("Directory/File Tree Begins -->", level=2)
        write_tree_docx(args.repo_path, doc, args, args.output_file, "", is_last=True, is_root=True)
        doc.add_heading("<-- Directory/File Tree Ends", level=2)
        doc.add_heading("File Content Begins -->", level=2)
        write_file_contents_in_order_docx(args.repo_path, doc, args)
        doc.add_heading("<-- File Content Ends", level=2)
        doc.save(args.output_file)
    else:
        with open(args.output_file, 'w', encoding='utf-8') as output_file:
            output_file.write("Repository Documentation\n")
            output_file.write(
            "This document provides a comprehensive overview of the repository's structure and contents.\n"
            "The first section, titled 'Directory/File Tree', displays the repository's hierarchy in a tree format.\n"
            "In this section, directories and files are listed using tree branches to indicate their structure and relationships.\n"
            "Following the tree representation, the 'File Content' section details the contents of each file in the repository.\n"
            "Each file's content is introduced with a '[File Begins]' marker followed by the file's relative path,\n"
            "and the content is displayed verbatim. The end of each file's content is marked with a '[File Ends]' marker.\n"
            "This format ensures a clear and orderly presentation of both the structure and the detailed contents of the repository.\n\n"
            )

            output_file.write("Directory/File Tree Begins -->\n\n")
            write_tree(args.repo_path, output_file, args, "", is_last=True, is_root=True)
            output_file.write("\n<-- Directory/File Tree Ends")
            output_file.write("\n\nFile Content Begin -->\n")
            write_file_contents_in_order(args.repo_path, output_file, args)
            output_file.write("\n<-- File Content Ends\n\n")
if __name__ == "__main__":
    main()
